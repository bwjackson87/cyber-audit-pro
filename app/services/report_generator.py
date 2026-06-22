"""
Report generation service.
Produces DOCX and PDF deliverables from project data.
"""

import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from sqlalchemy.orm import Session

from app.config import settings
from app.models import Project, Finding, AssessmentResponse, ReportExport
from app.services.scoring import compute_project_score
from app.services.template_manager import get_all_controls

log = logging.getLogger(__name__)

SEVERITY_COLORS = {
    "critical": (0xC0, 0x00, 0x00),
    "high":     (0xFF, 0x45, 0x00),
    "medium":   (0xFF, 0xA5, 0x00),
    "low":      (0x00, 0x70, 0xC0),
    "informational": (0x70, 0x70, 0x70),
}

SEVERITY_ORDER = ["critical", "high", "medium", "low", "informational"]
EFFORT_LABELS = {"low": "Low (< 1 week)", "medium": "Medium (1–4 weeks)", "high": "High (> 1 month)"}
STATUS_LABELS = {
    "compliant": "Compliant",
    "partial": "Partially Compliant",
    "non_compliant": "Non-Compliant",
    "na": "Not Applicable",
    "not_assessed": "Not Assessed",
}


def _safe_filename(name: str) -> str:
    return re.sub(r'[<>:"/\\|?*]', "_", name)


def _get_project_data(db: Session, project_id: int) -> dict:
    """Gather all data needed for any report type."""
    project: Project = db.query(Project).filter_by(id=project_id).first()
    if not project:
        raise ValueError(f"Project {project_id} not found")

    findings: list[Finding] = (
        db.query(Finding)
        .filter_by(project_id=project_id)
        .order_by(Finding.finding_ref)
        .all()
    )
    responses: list[AssessmentResponse] = (
        db.query(AssessmentResponse).filter_by(project_id=project_id).all()
    )
    score_data = compute_project_score(responses)

    controls = []
    if project.template and project.template.template_data:
        controls = get_all_controls(project.template.template_data)

    response_map = {r.control_id: r for r in responses}

    return {
        "project": project,
        "findings": findings,
        "responses": responses,
        "score_data": score_data,
        "controls": controls,
        "response_map": response_map,
        "generated_at": datetime.utcnow(),
    }


# ─────────────────────────────────────────────────────────────────────────────
#  DOCX generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_docx(db: Session, project_id: int, report_type: str = "full") -> Path:
    """
    Generate a DOCX report.  report_type: full | executive | findings | risk_register | remediation
    Returns the file path.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    data = _get_project_data(db, project_id)
    project: Project = data["project"]
    findings: list[Finding] = data["findings"]
    score_data: dict = data["score_data"]

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Styles ────────────────────────────────────────────────────────────────
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Calibri"
    style_normal.font.size = Pt(10)

    def heading(text: str, level: int = 1, color: tuple = (0x1E, 0x3A, 0x5F)):
        p = doc.add_heading(text, level=level)
        run = p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"
        return p

    def para(text: str = "", bold: bool = False, italic: bool = False, size: int = 10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = "Calibri"
        return p

    def add_labeled_para(label: str, value: str):
        p = doc.add_paragraph()
        r1 = p.add_run(f"{label}: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(value or "—")
        r2.font.size = Pt(10)

    def severity_cell_color(cell, severity: str):
        r, g, b = SEVERITY_COLORS.get(severity.lower(), (0x70, 0x70, 0x70))
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), f"{r:02X}{g:02X}{b:02X}")
        tcPr.append(shd)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.bold = True

    # ── Cover Page ───────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(project.consultant_company or "Cybersecurity Assessment")
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    r.font.name = "Calibri"

    doc.add_paragraph()
    t2 = doc.add_paragraph()
    t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    report_titles = {
        "full": "Cybersecurity Assessment Report",
        "executive": "Executive Summary",
        "findings": "Technical Findings Report",
        "risk_register": "Risk Register",
        "remediation": "Remediation Roadmap",
    }
    r2 = t2.add_run(report_titles.get(report_type, "Cybersecurity Assessment Report"))
    r2.font.size = Pt(24)
    r2.font.bold = True
    r2.font.name = "Calibri"

    doc.add_paragraph()
    t3 = doc.add_paragraph()
    t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = t3.add_run(f"Prepared for: {project.client.name}")
    r3.font.size = Pt(14)
    r3.font.name = "Calibri"

    doc.add_paragraph()
    info_table = doc.add_table(rows=5, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_rows = [
        ("Assessment Date", project.assessment_date.strftime("%B %d, %Y") if project.assessment_date else ""),
        ("Prepared By", project.consultant_name or ""),
        ("Title", project.consultant_title or ""),
        ("Organization", project.consultant_company or ""),
        ("Classification", "CONFIDENTIAL"),
    ]
    for i, (lbl, val) in enumerate(info_rows):
        row = info_table.rows[i]
        row.cells[0].text = lbl
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ── Disclaimer ────────────────────────────────────────────────────────────
    heading("Important Notice", level=1)
    disclaimer_text = project.report_disclaimer or settings.DEFAULT_DISCLAIMER
    p = doc.add_paragraph(disclaimer_text)
    p.runs[0].italic = True
    p.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # ── Executive Summary ─────────────────────────────────────────────────────
    heading("Executive Summary", level=1)

    if project.executive_summary:
        doc.add_paragraph(project.executive_summary)
        doc.add_paragraph()

    # Scorecard table
    heading("Assessment Scorecard", level=2)
    score_table = doc.add_table(rows=6, cols=2)
    score_table.style = "Table Grid"
    score_rows = [
        ("Overall Compliance Score", f"{score_data['overall_score']:.1f}%"),
        ("Maturity Level", score_data["maturity_level"]),
        ("Controls Assessed", str(score_data["counts"]["compliant"] + score_data["counts"]["partial"] + score_data["counts"]["non_compliant"])),
        ("Compliant", str(score_data["counts"]["compliant"])),
        ("Partially Compliant", str(score_data["counts"]["partial"])),
        ("Non-Compliant", str(score_data["counts"]["non_compliant"])),
    ]
    for i, (lbl, val) in enumerate(score_rows):
        row = score_table.rows[i]
        row.cells[0].text = lbl
        row.cells[1].text = val
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # Finding summary by severity
    heading("Finding Summary", level=2)
    sev_counts = {}
    for f in findings:
        s = f.severity or "informational"
        sev_counts[s] = sev_counts.get(s, 0) + 1

    if sev_counts:
        sev_table = doc.add_table(rows=1 + len(SEVERITY_ORDER), cols=2)
        sev_table.style = "Table Grid"
        sev_table.rows[0].cells[0].text = "Severity"
        sev_table.rows[0].cells[1].text = "Count"
        for cell in sev_table.rows[0].cells:
            cell.paragraphs[0].runs[0].bold = True
        for i, sev in enumerate(SEVERITY_ORDER, 1):
            row = sev_table.rows[i]
            row.cells[0].text = sev.title()
            row.cells[1].text = str(sev_counts.get(sev, 0))
    else:
        doc.add_paragraph("No findings recorded.")

    doc.add_paragraph()

    # Scope and methodology
    heading("Scope", level=2)
    doc.add_paragraph(project.scope or "Not specified.")
    if project.out_of_scope:
        heading("Out of Scope", level=2)
        doc.add_paragraph(project.out_of_scope)
    if project.methodology:
        heading("Methodology", level=2)
        doc.add_paragraph(project.methodology)
    if project.assumptions:
        heading("Assumptions and Limitations", level=2)
        doc.add_paragraph(project.assumptions)
        if project.limitations:
            doc.add_paragraph(project.limitations)

    # Stop here if executive only
    if report_type == "executive":
        doc.add_page_break()
        _add_docx_footer_page(doc)
        return _save_docx(doc, db, project, report_type)

    doc.add_page_break()

    # ── Detailed Findings ─────────────────────────────────────────────────────
    if report_type in ("full", "findings"):
        heading("Detailed Findings", level=1)

        confirmed = [f for f in findings if f.is_confirmed]
        observations = [f for f in findings if not f.is_confirmed]

        if not confirmed and not observations:
            doc.add_paragraph("No findings recorded for this assessment.")
        else:
            for sev in SEVERITY_ORDER:
                sev_findings = [f for f in confirmed if (f.severity or "").lower() == sev]
                if not sev_findings:
                    continue
                heading(f"{sev.title()} Findings", level=2)
                for finding in sev_findings:
                    _add_finding_to_docx(doc, finding, heading, para, add_labeled_para, severity_cell_color)

            if observations:
                heading("Observations", level=2)
                para("Observations are noted areas of concern that do not rise to the level of a formal finding at this time.", italic=True)
                for finding in observations:
                    _add_finding_to_docx(doc, finding, heading, para, add_labeled_para, severity_cell_color)

        doc.add_page_break()

    # ── Risk Register ─────────────────────────────────────────────────────────
    if report_type in ("full", "risk_register"):
        heading("Risk Register", level=1)
        para("Risk Score = Likelihood (1–5) × Impact (1–5). Maximum score = 25.", italic=True)
        doc.add_paragraph()

        if findings:
            rr_table = doc.add_table(rows=1 + len(findings), cols=7)
            rr_table.style = "Table Grid"
            headers = ["Ref", "Title", "Severity", "Likelihood", "Impact", "Risk Score", "Priority"]
            for i, h in enumerate(headers):
                cell = rr_table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True

            for i, finding in enumerate(findings, 1):
                row = rr_table.rows[i]
                row.cells[0].text = finding.finding_ref or ""
                row.cells[1].text = finding.title or ""
                row.cells[2].text = (finding.severity or "").title()
                row.cells[3].text = str(finding.likelihood or "")
                row.cells[4].text = str(finding.impact or "")
                row.cells[5].text = str(int(finding.risk_score or 0))
                row.cells[6].text = str(finding.remediation_priority or "")
                if finding.severity:
                    severity_cell_color(row.cells[2], finding.severity)
        else:
            doc.add_paragraph("No findings recorded.")

        doc.add_page_break()

    # ── Remediation Roadmap ───────────────────────────────────────────────────
    if report_type in ("full", "remediation"):
        heading("Remediation Roadmap", level=1)
        para("Items are ordered by remediation priority (1 = immediate, 5 = long-term).", italic=True)
        doc.add_paragraph()

        sorted_findings = sorted(
            [f for f in findings if f.remediation_priority],
            key=lambda x: (x.remediation_priority or 99, x.finding_ref or "")
        )

        for priority in range(1, 6):
            pf = [f for f in sorted_findings if f.remediation_priority == priority]
            if not pf:
                continue
            priority_labels = {
                1: "Priority 1 — Immediate (< 30 days)",
                2: "Priority 2 — Short-Term (30–90 days)",
                3: "Priority 3 — Medium-Term (90–180 days)",
                4: "Priority 4 — Long-Term (180+ days)",
                5: "Priority 5 — Strategic / Ongoing",
            }
            heading(priority_labels.get(priority, f"Priority {priority}"), level=2)
            for f in pf:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(f"[{f.finding_ref}] {f.title}")
                r.bold = True
                r.font.size = Pt(10)
                if f.recommendation:
                    p2 = doc.add_paragraph(f.recommendation)
                    p2.paragraph_format.left_indent = Inches(0.25)
                    p2.runs[0].font.size = Pt(9)
                add_labeled_para(
                    "Effort",
                    EFFORT_LABELS.get(f.remediation_effort or "", f.remediation_effort or "—")
                )

        doc.add_page_break()

    _add_docx_footer_page(doc)
    return _save_docx(doc, db, project, report_type)


def _add_finding_to_docx(doc, finding, heading_fn, para_fn, labeled_fn, sev_color_fn):
    from docx.shared import Pt, RGBColor
    sev = (finding.severity or "informational").lower()
    r, g, b = SEVERITY_COLORS.get(sev, (0x70, 0x70, 0x70))

    p = doc.add_heading(f"{finding.finding_ref or 'F-??'}: {finding.title}", level=3)
    if p.runs:
        p.runs[0].font.color.rgb = RGBColor(r, g, b)

    labeled_fn("Severity", (finding.severity or "").title())
    labeled_fn("Control Reference", finding.control_ref or "—")
    labeled_fn("Affected Systems", finding.affected_systems or "—")

    if finding.description:
        heading_fn("Description", level=4)
        doc.add_paragraph(finding.description)
    if finding.current_state:
        labeled_fn("Current State", finding.current_state)
    if finding.required_state:
        labeled_fn("Required State", finding.required_state)
    if finding.risk_impact:
        heading_fn("Risk Impact", level=4)
        doc.add_paragraph(finding.risk_impact)
    if finding.recommendation:
        heading_fn("Recommendation", level=4)
        doc.add_paragraph(finding.recommendation)
    if finding.remediation_detail:
        heading_fn("Remediation Detail", level=4)
        doc.add_paragraph(finding.remediation_detail)
    if finding.compensating_controls:
        labeled_fn("Compensating Controls", finding.compensating_controls)

    labeled_fn("Remediation Effort", EFFORT_LABELS.get(finding.remediation_effort or "", "—"))
    labeled_fn("Priority", str(finding.remediation_priority or "—"))

    doc.add_paragraph()


def _add_docx_footer_page(doc):
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph(f"Report generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(8)
    p.runs[0].italic = True


def _save_docx(doc, db: Session, project: Project, report_type: str) -> Path:
    settings.EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(
        f"{project.client.name}_{project.name}_{report_type}_{datetime.utcnow().strftime('%Y%m%d')}.docx"
    )
    path = settings.EXPORTS_DIR / filename
    doc.save(str(path))

    export = ReportExport(
        project_id=project.id,
        report_type=report_type,
        format="docx",
        filepath=str(path),
    )
    db.add(export)
    db.commit()
    log.info("DOCX report saved: %s", path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
#  PDF generation
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf(db: Session, project_id: int, report_type: str = "full") -> Path:
    """Generate a PDF report using ReportLab."""
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    data = _get_project_data(db, project_id)
    project: Project = data["project"]
    findings: list[Finding] = data["findings"]
    score_data: dict = data["score_data"]

    settings.EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(
        f"{project.client.name}_{project.name}_{report_type}_{datetime.utcnow().strftime('%Y%m%d')}.pdf"
    )
    path = settings.EXPORTS_DIR / filename

    styles = getSampleStyleSheet()
    navy = colors.HexColor("#1E3A5F")
    dark = colors.HexColor("#1a1a1a")

    style_title = ParagraphStyle("Title2", parent=styles["Title"], textColor=navy, fontSize=24, spaceAfter=12)
    style_h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=navy, fontSize=16, spaceBefore=16, spaceAfter=6)
    style_h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=navy, fontSize=13, spaceBefore=12, spaceAfter=4)
    style_h3 = ParagraphStyle("H3", parent=styles["Heading3"], textColor=dark, fontSize=11, spaceBefore=8, spaceAfter=4)
    style_body = ParagraphStyle("Body2", parent=styles["Normal"], fontSize=10, spaceAfter=6)
    style_small = ParagraphStyle("Small", parent=styles["Normal"], fontSize=8, textColor=colors.grey, spaceAfter=4)
    style_center = ParagraphStyle("Center", parent=styles["Normal"], alignment=TA_CENTER, fontSize=10)
    style_label = ParagraphStyle("Label", parent=styles["Normal"], fontSize=10, textColor=navy, fontName="Helvetica-Bold")

    sev_color_map = {
        "critical": colors.HexColor("#C00000"),
        "high": colors.HexColor("#FF4500"),
        "medium": colors.HexColor("#FFA500"),
        "low": colors.HexColor("#0070C0"),
        "informational": colors.HexColor("#707070"),
    }

    story = []

    report_titles = {
        "full": "Cybersecurity Assessment Report",
        "executive": "Executive Summary",
        "findings": "Technical Findings Report",
        "risk_register": "Risk Register",
        "remediation": "Remediation Roadmap",
    }

    # Cover
    story.append(Spacer(1, 1.5 * inch))
    story.append(Paragraph(project.consultant_company or "Cybersecurity Assessment", style_h1))
    story.append(Paragraph(report_titles.get(report_type, "Cybersecurity Assessment Report"), style_title))
    story.append(Spacer(1, 0.3 * inch))
    story.append(Paragraph(f"Prepared for: <b>{project.client.name}</b>", style_body))
    story.append(Spacer(1, 0.5 * inch))

    cover_data = [
        ["Assessment Date", project.assessment_date.strftime("%B %d, %Y") if project.assessment_date else ""],
        ["Prepared By", project.consultant_name or ""],
        ["Title", project.consultant_title or ""],
        ["Organization", project.consultant_company or ""],
        ["Classification", "CONFIDENTIAL"],
    ]
    cover_table = Table(cover_data, colWidths=[2 * inch, 4 * inch])
    cover_table.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("TEXTCOLOR", (0, 0), (0, -1), navy),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1), [colors.HexColor("#F5F7FA"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(cover_table)
    story.append(PageBreak())

    # Disclaimer
    story.append(Paragraph("Important Notice", style_h1))
    disclaimer_text = project.report_disclaimer or settings.DEFAULT_DISCLAIMER
    story.append(Paragraph(f"<i>{disclaimer_text}</i>", style_small))
    story.append(Spacer(1, 0.3 * inch))

    # Executive Summary
    story.append(Paragraph("Executive Summary", style_h1))
    if project.executive_summary:
        story.append(Paragraph(project.executive_summary, style_body))
        story.append(Spacer(1, 0.2 * inch))

    # Scorecard
    story.append(Paragraph("Assessment Scorecard", style_h2))
    score_table_data = [
        ["Metric", "Value"],
        ["Overall Compliance Score", f"{score_data['overall_score']:.1f}%"],
        ["Maturity Level", score_data["maturity_level"]],
        ["Controls Compliant", str(score_data["counts"]["compliant"])],
        ["Partially Compliant", str(score_data["counts"]["partial"])],
        ["Non-Compliant", str(score_data["counts"]["non_compliant"])],
        ["Not Applicable", str(score_data["counts"]["na"])],
    ]
    st = Table(score_table_data, colWidths=[3 * inch, 3 * inch])
    st.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), navy),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 10),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F5F7FA"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(st)
    story.append(Spacer(1, 0.2 * inch))

    # Finding counts by severity
    story.append(Paragraph("Finding Summary", style_h2))
    sev_counts = {}
    for f in findings:
        s = f.severity or "informational"
        sev_counts[s] = sev_counts.get(s, 0) + 1

    if sev_counts:
        sev_data = [["Severity", "Count"]]
        for sev in SEVERITY_ORDER:
            sev_data.append([sev.title(), str(sev_counts.get(sev, 0))])
        st2 = Table(sev_data, colWidths=[3 * inch, 1.5 * inch])
        sev_styles = [
            ("BACKGROUND", (0, 0), (-1, 0), navy),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
            ("PADDING", (0, 0), (-1, -1), 6),
        ]
        for i, sev in enumerate(SEVERITY_ORDER, 1):
            c = sev_color_map.get(sev, colors.grey)
            sev_styles.append(("BACKGROUND", (0, i), (0, i), c))
            sev_styles.append(("TEXTCOLOR", (0, i), (0, i), colors.white))
            sev_styles.append(("FONTNAME", (0, i), (0, i), "Helvetica-Bold"))
        st2.setStyle(TableStyle(sev_styles))
        story.append(st2)
    else:
        story.append(Paragraph("No findings recorded.", style_body))

    story.append(Spacer(1, 0.2 * inch))

    # Scope
    story.append(Paragraph("Scope", style_h2))
    story.append(Paragraph(project.scope or "Not specified.", style_body))
    if project.out_of_scope:
        story.append(Paragraph("Out of Scope", style_h2))
        story.append(Paragraph(project.out_of_scope, style_body))
    if project.methodology:
        story.append(Paragraph("Methodology", style_h2))
        story.append(Paragraph(project.methodology, style_body))
    if project.assumptions:
        story.append(Paragraph("Assumptions and Limitations", style_h2))
        story.append(Paragraph(project.assumptions, style_body))
        if project.limitations:
            story.append(Paragraph(project.limitations, style_body))

    if report_type == "executive":
        story.append(_pdf_footer_para(style_small))
        doc_obj = SimpleDocTemplate(str(path), pagesize=LETTER,
                                    topMargin=inch, bottomMargin=inch,
                                    leftMargin=1.25 * inch, rightMargin=1.25 * inch)
        doc_obj.build(story)
        return _record_export(db, project, report_type, "pdf", path)

    story.append(PageBreak())

    # Detailed Findings
    if report_type in ("full", "findings"):
        story.append(Paragraph("Detailed Findings", style_h1))
        confirmed = [f for f in findings if f.is_confirmed]
        observations = [f for f in findings if not f.is_confirmed]

        if not confirmed and not observations:
            story.append(Paragraph("No findings recorded.", style_body))
        else:
            for sev in SEVERITY_ORDER:
                sev_findings = [f for f in confirmed if (f.severity or "").lower() == sev]
                if not sev_findings:
                    continue
                story.append(Paragraph(f"{sev.title()} Findings", style_h2))
                for finding in sev_findings:
                    _add_finding_to_pdf(story, finding, style_h3, style_body, style_label,
                                        style_small, sev_color_map)

            if observations:
                story.append(Paragraph("Observations", style_h2))
                story.append(Paragraph(
                    "<i>Observations are noted areas of concern that do not rise to the level of a formal finding.</i>",
                    style_small))
                for finding in observations:
                    _add_finding_to_pdf(story, finding, style_h3, style_body, style_label,
                                        style_small, sev_color_map)

        story.append(PageBreak())

    # Risk Register
    if report_type in ("full", "risk_register"):
        story.append(Paragraph("Risk Register", style_h1))
        story.append(Paragraph("<i>Risk Score = Likelihood (1–5) × Impact (1–5)</i>", style_small))
        story.append(Spacer(1, 0.15 * inch))

        if findings:
            rr_data = [["Ref", "Title", "Severity", "L", "I", "Score", "Priority"]]
            for f in findings:
                rr_data.append([
                    f.finding_ref or "",
                    Paragraph(f.title or "", style_small),
                    (f.severity or "").title(),
                    str(f.likelihood or ""),
                    str(f.impact or ""),
                    str(int(f.risk_score or 0)),
                    str(f.remediation_priority or ""),
                ])
            rr_table = Table(rr_data, colWidths=[0.6*inch, 2.8*inch, 0.9*inch, 0.3*inch, 0.3*inch, 0.5*inch, 0.5*inch])
            rr_styles = [
                ("BACKGROUND", (0, 0), (-1, 0), navy),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F5F7FA"), colors.white]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("PADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
            for i, f in enumerate(findings, 1):
                c = sev_color_map.get((f.severity or "").lower(), colors.grey)
                rr_styles.append(("BACKGROUND", (2, i), (2, i), c))
                rr_styles.append(("TEXTCOLOR", (2, i), (2, i), colors.white))
            rr_table.setStyle(TableStyle(rr_styles))
            story.append(rr_table)
        else:
            story.append(Paragraph("No findings recorded.", style_body))

        story.append(PageBreak())

    # Remediation Roadmap
    if report_type in ("full", "remediation"):
        story.append(Paragraph("Remediation Roadmap", style_h1))
        story.append(Paragraph(
            "<i>Items ordered by priority: 1 = Immediate action required, 5 = Long-term/strategic initiative.</i>",
            style_small))
        story.append(Spacer(1, 0.15 * inch))

        sorted_findings = sorted(
            [f for f in findings if f.remediation_priority],
            key=lambda x: (x.remediation_priority or 99, x.finding_ref or "")
        )
        priority_labels = {
            1: "Priority 1 — Immediate (< 30 days)",
            2: "Priority 2 — Short-Term (30–90 days)",
            3: "Priority 3 — Medium-Term (90–180 days)",
            4: "Priority 4 — Long-Term (180+ days)",
            5: "Priority 5 — Strategic / Ongoing",
        }
        for priority in range(1, 6):
            pf = [f for f in sorted_findings if f.remediation_priority == priority]
            if not pf:
                continue
            story.append(Paragraph(priority_labels[priority], style_h2))
            rm_data = [["Ref", "Title", "Effort", "Recommendation"]]
            for f in pf:
                rm_data.append([
                    f.finding_ref or "",
                    Paragraph(f.title or "", style_small),
                    EFFORT_LABELS.get(f.remediation_effort or "", "—"),
                    Paragraph(f.recommendation or "—", style_small),
                ])
            rm_table = Table(rm_data, colWidths=[0.6*inch, 2*inch, 1.2*inch, 2.6*inch])
            rm_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), navy),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.HexColor("#F5F7FA"), colors.white]),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#CCCCCC")),
                ("PADDING", (0, 0), (-1, -1), 4),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            story.append(rm_table)
            story.append(Spacer(1, 0.2 * inch))

    story.append(_pdf_footer_para(style_small))

    doc_obj = SimpleDocTemplate(str(path), pagesize=LETTER,
                                topMargin=inch, bottomMargin=inch,
                                leftMargin=1.25 * inch, rightMargin=1.25 * inch)
    doc_obj.build(story)
    return _record_export(db, project, report_type, "pdf", path)


def _add_finding_to_pdf(story, finding, style_h3, style_body, style_label, style_small, sev_color_map):
    from reportlab.lib import colors
    from reportlab.platypus import Spacer, HRFlowable
    from reportlab.lib.units import inch

    sev = (finding.severity or "informational").lower()
    c = sev_color_map.get(sev, colors.grey)

    story.append(Paragraph(f"{finding.finding_ref or 'F-??'}: {finding.title}", style_h3))

    meta_data = [
        ["Severity", (finding.severity or "").title()],
        ["Control Reference", finding.control_ref or "—"],
        ["Affected Systems", finding.affected_systems or "—"],
    ]
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.units import inch
    mt = Table(meta_data, colWidths=[1.5 * inch, 5 * inch])
    mt.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1E3A5F")),
        ("PADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(mt)

    if finding.description:
        story.append(Paragraph("<b>Description</b>", style_label))
        story.append(Paragraph(finding.description, style_body))
    if finding.current_state:
        story.append(Paragraph(f"<b>Current State:</b> {finding.current_state}", style_body))
    if finding.required_state:
        story.append(Paragraph(f"<b>Required State:</b> {finding.required_state}", style_body))
    if finding.risk_impact:
        story.append(Paragraph("<b>Risk Impact</b>", style_label))
        story.append(Paragraph(finding.risk_impact, style_body))
    if finding.recommendation:
        story.append(Paragraph("<b>Recommendation</b>", style_label))
        story.append(Paragraph(finding.recommendation, style_body))
    if finding.remediation_detail:
        story.append(Paragraph("<b>Remediation Detail</b>", style_label))
        story.append(Paragraph(finding.remediation_detail, style_body))
    if finding.compensating_controls:
        story.append(Paragraph(f"<b>Compensating Controls:</b> {finding.compensating_controls}", style_small))

    story.append(Paragraph(
        f"<b>Effort:</b> {EFFORT_LABELS.get(finding.remediation_effort or '', '—')}  |  "
        f"<b>Priority:</b> {finding.remediation_priority or '—'}",
        style_small))
    story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#CCCCCC"), spaceAfter=8))


def _pdf_footer_para(style):
    from reportlab.platypus import Paragraph
    return Paragraph(
        f"<i>Report generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')} — CONFIDENTIAL</i>",
        style
    )


def _record_export(db: Session, project: Project, report_type: str, fmt: str, path: Path) -> Path:
    export = ReportExport(
        project_id=project.id,
        report_type=report_type,
        format=fmt,
        filepath=str(path),
    )
    db.add(export)
    db.commit()
    log.info("PDF report saved: %s", path)
    return path


# ─────────────────────────────────────────────────────────────────────────────
#  CSV export
# ─────────────────────────────────────────────────────────────────────────────

def generate_csv_risk_register(db: Session, project_id: int) -> Path:
    import csv
    data = _get_project_data(db, project_id)
    project = data["project"]
    findings = data["findings"]

    settings.EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    filename = _safe_filename(
        f"{project.client.name}_{project.name}_risk_register_{datetime.utcnow().strftime('%Y%m%d')}.csv"
    )
    path = settings.EXPORTS_DIR / filename

    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow([
            "Finding Ref", "Type", "Control Ref", "Title", "Severity",
            "Likelihood", "Impact", "Risk Score", "Affected Systems",
            "Recommendation", "Remediation Effort", "Priority", "Confirmed"
        ])
        for fi in findings:
            writer.writerow([
                fi.finding_ref or "",
                fi.finding_type or "",
                fi.control_ref or "",
                fi.title or "",
                fi.severity or "",
                fi.likelihood or "",
                fi.impact or "",
                fi.risk_score or "",
                fi.affected_systems or "",
                fi.recommendation or "",
                fi.remediation_effort or "",
                fi.remediation_priority or "",
                "Yes" if fi.is_confirmed else "No",
            ])

    export = ReportExport(
        project_id=project.id,
        report_type="risk_register",
        format="csv",
        filepath=str(path),
    )
    db.add(export)
    db.commit()
    return path
